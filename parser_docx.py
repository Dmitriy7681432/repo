# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
def page_break():
    """Метод Document.add_page_break()"""
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph()
    return p

# https://docs-python.ru/packages/modul-python-docx-python/obekt-table/
doc  = Document()

items = (
    ('Время', 'Активность', 'Примечание'),
    ('Начало работы - 8:50', '', ''),
    ('8:50-9:00', 'Зарядка для глаз, отвлечение от компьютера', 'В соответствии с правилами техники безопасности'),
    ('9:00-10:00', '', ''),
    ('10:00-10:15', 'Официальный перерыв', ''),
    ('10:15-11:05', '', ''),
    ('11:05-11:15', 'Зарядка для глаз, отвлечение от компьютера', 'В соответствии с правилами техники безопасности'),
    ('11:15-12:00', '', ''),
    ('12:00-13:00', 'Обед', ''),
    ('13:00-13:50', '', ''),
    ('13:50-14:00', 'Зарядка для глаз, отвлечение от компьютера', 'В соответствии с правилами техники безопасности'),
    ('14:00-15:00', '', ''),
    ('15:00-15:15', 'Официальный перерыв', ''),
    ('15:15-16:05', '', ''),
    ('16:05-16:15', 'Зарядка для глаз, отвлечение от компьютера', 'В соответствии с правилами техники безопасности'),
    ('16:15 - Окончание работы', 'Заполнение ежедневного отчета о проделанной работе', ''),
)
# добавляем пустую таблицу 2х2 ячейки
doc.add_paragraph('ФИО: Алексин Д.В.'+'\n'+'Отчетный месяц: март 2023 года')

item = ' 1 марта 2023 года'

table = doc.add_table(1, 1)
table.style = 'Table Grid'
head_cells = table.rows[0].cells
p = head_cells[0].paragraphs[0]
# название колонки
p.add_run(item)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# for row in items:
#     for i, item in enumerate(row):
#         print(str(item))
table = doc.add_table(0, 3)
table.style = 'Table Grid'
for row in items:
    # добавляем строку с ячейками к объекту таблицы
    cells = table.add_row().cells
    for i, item in enumerate(row):
        # вставляем данные в ячейки
        cells[i].text = str(item)
        print(cells[i].text)
        # если последняя ячейка
        if i == 2:
            # изменим шрифт
            cells[i].paragraphs[0].runs[0].font.name = 'Arial'

# cells = table.add_row().cells
# cells[0].text = str(item)
Cell = table.cell(1,1)
Cell.width
page_break()

#--------------------------------------------------------------------------------------
items = (
    (7, '1024', 'Плюшевые котята'),
    (3, '2042', 'Меховые пчелы'),
    (1, '1288', 'Ошейники для пуделей'),
)
# добавляем таблицу с одной строкой
# для заполнения названий колонок
table = doc.add_table(1, len(items[0]))
table.style = 'Table Grid'
# Получаем строку с колонками из добавленной таблицы
head_cells = table.rows[0].cells
# for i in head_cells:
#     print(i.text)
for i, item in enumerate(['Кол-во', 'ID', 'Описание']):
    p = head_cells[i].paragraphs[0]
    # название колонки
    p.add_run(item).bold = True
    # выравниваем посередине
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # добавляем данные к существующей таблице
for row in items:
    # добавляем строку с ячейками к объекту таблицы
    cells = table.add_row().cells
    for i, item in enumerate(row):
        # вставляем данные в ячейки
        cells[i].text = str(item)
        # если последняя ячейка
        if i == 2:
            # изменим шрифт
            cells[i].paragraphs[0].runs[0].font.name = 'Arial'
doc.save('test1.docx')
# #--------------------------------------------------------------------------------------



# table = doc.add_table(rows=2, cols=2)
# cell = table.cell(1, 1)
# cell.text = 'Бык'
# rc = cell.paragraphs[0].runs[0]

# row = table.rows[1]
# # запись данных в ячейки
# row.cells[0].text = 'Заяц'
# row.cells[1].text = 'Волк'

# rc.font.name = 'Arial'
# rc.font.bold = True
doc.save('test.docx')
